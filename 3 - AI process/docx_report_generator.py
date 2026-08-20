import os
import requests
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def generate_docx_report(enkf_filter, census_summary, analysis_dir="report", plots_dir="report", script_dir="."):
    """
    Generates an exhaustive Microsoft Word (.docx) technical document inside analysis_dir (default 'report/')
    containing all user commands, hotkey tables, module directives, and mathematical formulations.
    """
    os.makedirs(analysis_dir, exist_ok=True)
    docx_path = os.path.abspath(os.path.join(analysis_dir, "ecological_analysis_report.docx"))
    
    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)
        
    # Document Header Title
    p_title = doc.add_paragraph()
    r_t = p_title.add_run("AquaPulse: Commands, Directives & Mathematical Report")
    r_t.font.name = 'Helvetica'
    r_t.font.size = Pt(22)
    r_t.font.bold = True
    r_t.font.color.rgb = RGBColor(15, 23, 42)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(14)
    r_s = p_sub.add_run("Exhaustive Compilation of Interactive Commands, Architecture Directives & Mathematical Formulas")
    r_s.font.name = 'Helvetica'
    r_s.font.size = Pt(11)
    r_s.font.color.rgb = RGBColor(37, 99, 235)

    # --- SECTION 1: INTERACTIVE KEYBOARD COMMANDS & HOTKEYS ---
    h1 = doc.add_heading("1. Interactive Keyboard Commands & Hotkey Registry", level=1)
    h1.runs[0].font.color.rgb = RGBColor(37, 99, 235)
    
    doc.add_paragraph(
        "Below is the complete registry of interactive hotkey commands and controls available in the AquaPulse dashboard:"
    )

    t_cmd = doc.add_table(rows=1, cols=3)
    t_cmd.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, text in enumerate(["Hotkey / Keybind", "Command Action", "Target Subsystem"]):
        cell = t_cmd.rows[0].cells[idx]
        cell.text = text
        set_cell_background(cell, "0F172A")
        for r in cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.bold = True

    commands = [
        ("[S]", "Cycle Target Filter (ALL TARGETS, FISH ONLY, DIVERS & OCTOPUS)", "Object Detection Filter"),
        ("[F]", "Cycle Vision FX Modes (NORMAL, THERMAL, NIGHT VISION, SONAR EDGE, CLAHE)", "Image Processing FX"),
        ("[V]", "Toggle Kinematic Motion Vector Overlay Arrows", "Kinematic Vectors"),
        ("[H]", "Toggle Frame Accumulation Density Heatmap Overlay", "Spatial Density Heatmap"),
        ("[G]", "Toggle Sonar HUD Grid Overlay", "Spatial HUD Grid"),
        ("[Z]", "Toggle Target Magnifier 2.0x PiP Zoom Window", "Optical Zoom Magnifier"),
        ("[A]", "Toggle Underwater Adaptive CLAHE Contrast Enhancement", "Contrast Boost"),
        ("[W]", "Toggle Water Ripple Video Composite Overlay", "VFX Composite"),
        ("[M]", "Toggle Statistical Analyzer & 50-Step Future Predictor Mode", "EnKF Predictive Engine"),
        ("[L]", "Toggle Telemetry Audio Language Mode (EN / DE)", "Audio & UI Localizer"),
        ("[C]", "Trigger Dr. Daniel Pauly Audio Link & GBIF Dialogue", "LLM Audio Link"),
        ("[J]", "Trigger Johnny Silverhand Relic Hologram / GIF Construct", "Cybernetic Relic"),
        ("[CTRL+T]", "Open Interactive User Chat Box for Ollama LLM Query", "User Dialogue Input"),
        ("[+] / [-]", "Adjust Object Detection Confidence Threshold (+/- 0.05)", "YOLO Threshold Filter"),
        ("[O]", "Open OS GUI File Dialog to Select Any Custom Input Video", "Video Engine Source"),
        ("[Q]", "Initiate System Exit and Automated CSV, PNG Plot, PDF & DOCX Export", "Export Manager")
    ]

    for r_idx, (k, a, s) in enumerate(commands, start=1):
        row = t_cmd.add_row().cells
        row[0].text = k
        row[1].text = a
        row[2].text = s
        bg = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c in row:
            set_cell_background(c, bg)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # --- SECTION 2: SYSTEM ARCHITECTURE DIRECTIVES & COMMANDS ---
    h2 = doc.add_heading("2. System Architecture Commands & Module Specifications", level=1)
    h2.runs[0].font.color.rgb = RGBColor(37, 99, 235)

    t_mod = doc.add_table(rows=1, cols=3)
    t_mod.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, text in enumerate(["Module Command", "Architectural Specification", "Output Target"]):
        cell = t_mod.rows[0].cells[idx]
        cell.text = text
        set_cell_background(cell, "1E293B")
        for r in cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.bold = True

    modules_info = [
        ("MODULE 1", "Deduplicated tracking via species_unique_ids = defaultdict(set) & fish_counts.csv with Biomass % and Section B Methodology.", "report/fish_counts.csv"),
        ("MODULE 2", "Stochastic SDEs via Euler-Maruyama discrete time steps for predator-prey dynamics with Gaussian noise.", "stochastic_enkf.py"),
        ("MODULE 3", "100-Member Ensemble Kalman Filter (EnKF) for state forecast, covariance Pf, Kalman Gain K, and extinction risk prob.", "stochastic_enkf.py"),
        ("MODULE 4", "Matplotlib 3-figure diagnostic suite: Time-Series, 50-Step Monte Carlo Forecast, and Phase Space Orbit.", "report/*.png"),
        ("MODULE 5", "ReportLab PDF publication compiler with 4 structured sections and Ollama LLM integration.", "report/ecological_analysis_report.pdf"),
        ("MODULE 6", "Application integration into main YOLO loop, video pipeline, and automated shutdown export sequence.", "main.py")
    ]

    for r_idx, (m_cmd, m_spec, m_target) in enumerate(modules_info, start=1):
        row = t_mod.add_row().cells
        row[0].text = m_cmd
        row[1].text = m_spec
        row[2].text = m_target
        bg = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c in row:
            set_cell_background(c, bg)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # --- SECTION 3: ALL MATHEMATICAL FORMULAS ---
    h3 = doc.add_heading("3. Complete Mathematical Formulations & Equations", level=1)
    h3.runs[0].font.color.rgb = RGBColor(37, 99, 235)

    doc.add_paragraph("3.1. General Euler-Maruyama Discrete Integration Formula:").runs[0].bold = True
    p1 = doc.add_paragraph("Z_{n+1} = Z_n + Δt f(Z_n) + √(Δt) ζ_n,   where ζ_n ~ N(0, 1)")
    p1.runs[0].font.name = 'Courier New'

    doc.add_paragraph("3.2. Coupled Stochastic Lotka-Volterra Differential Equations:").runs[0].bold = True
    doc.add_paragraph("Prey Population (X_{n+1} - Tracked by YOLO):").runs[0].italic = True
    p2b = doc.add_paragraph("X_{n+1} = X_n + Δt (α X_n - β X_n Y_n) + √(Δt) σ_X X_n ζ_n^X")
    p2b.runs[0].font.name = 'Courier New'
    
    doc.add_paragraph("Predator Population (Y_{n+1} - Hidden Latent Variable):").runs[0].italic = True
    p2d = doc.add_paragraph("Y_{n+1} = Y_n + Δt (δ X_n Y_n - γ Y_n) + √(Δt) σ_Y Y_n ζ_n^Y")
    p2d.runs[0].font.name = 'Courier New'
    
    p2e = doc.add_paragraph("Parameters: α = 0.1, β = 0.02, γ = 0.1, δ = 0.01, σ_X = 0.05, σ_Y = 0.05, Δt = 0.1")
    p2e.runs[0].font.size = Pt(9)

    eq1_path = os.path.join(plots_dir, "latex_eq1.png")
    eq2_path = os.path.join(plots_dir, "latex_eq2.png")
    eq3_path = os.path.join(plots_dir, "latex_eq3.png")

    if os.path.exists(eq1_path):
        doc.add_picture(eq1_path, width=Inches(4.5))
    if os.path.exists(eq2_path):
        doc.add_picture(eq2_path, width=Inches(4.5))

    doc.add_paragraph("3.3. Ensemble Initialization (N = 100 Members):").runs[0].bold = True
    p3 = doc.add_paragraph("X_i ~ N(X_init, 1.0),   Y_i ~ N(10.0, 3.0^2),   for i = 1, 2, ..., 100")
    p3.runs[0].font.name = 'Courier New'

    doc.add_paragraph("3.4. Background Error Covariance Matrix (P^f):").runs[0].bold = True
    p4 = doc.add_paragraph("P^f = (1 / (N - 1)) Σ_{i=1}^N (x_i^f - x̄^f)(x_i^f - x̄^f)^T")
    p4.runs[0].font.name = 'Courier New'

    doc.add_paragraph("3.5. Kalman Gain Matrix (K) & Observation Matrix (H):").runs[0].bold = True
    p5 = doc.add_paragraph("H = [1, 0],   R = 4.0,   K = P^f H^T (H P^f H^T + R)^(-1)")
    p5.runs[0].font.name = 'Courier New'

    doc.add_paragraph("3.6. Ensemble State Analysis Update:").runs[0].bold = True
    p6 = doc.add_paragraph("x_i^a = x_i^f + K (z_k + ε_i - H x_i^f),   where ε_i ~ N(0, R)")
    p6.runs[0].font.name = 'Courier New'

    if os.path.exists(eq3_path):
        doc.add_picture(eq3_path, width=Inches(5.0))

    doc.add_paragraph("3.7. Extinction Risk Probability Formula (X_crit = 2.0):").runs[0].bold = True
    p7 = doc.add_paragraph("P(Extinction) = (1 / N) Σ_{i=1}^N I(X_i ≤ 2.0) * 100%")
    p7.runs[0].font.name = 'Courier New'

    doc.add_paragraph("3.8. Biomass Percentage Formula:").runs[0].bold = True
    p8 = doc.add_paragraph("Percentage of Total Biomass (%) = (Deduplicated Count(S_i) / Total Unique Count) * 100%")
    p8.runs[0].font.name = 'Courier New'

    doc.save(docx_path)
    print(f"[DOCX Generator] Successfully compiled Microsoft Word report: {docx_path}")
    return docx_path
