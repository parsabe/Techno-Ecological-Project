import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_system_architecture_docx(output_path):
    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_t = title_p.add_run("AquaPulse Aquatic Vision Tracking & Data Assimilation")
    run_t.bold = True
    run_t.font.size = Pt(22)
    run_t.font.color.rgb = RGBColor(0x00, 0x4B, 0x87)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(18)
    run_s = sub_p.add_run("Technical System Architecture & Mathematical Framework Documentation")
    run_s.font.size = Pt(14)
    run_s.italic = True
    run_s.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(24)
    run_m = meta_p.add_run("Author: AquaPulse AI Research Group  |  Department of Computer Vision & Data Assimilation\nDate: August 2026")
    run_m.font.size = Pt(10)
    run_m.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_heading("Abstract", level=1)
    abs_p = doc.add_paragraph(
        "This document provides a comprehensive technical reference for the AquaPulse AI Process codebase. "
        "The system combines real-time underwater computer vision (YOLOv8 fine-tuned models), multi-algorithm persistent "
        "target tracking (BoT-SORT and ByteTrack), deduplicated biological census enumeration, and a stochastic Ensemble "
        "Kalman Filter (EnKF) data assimilation framework (M=100 particles). Furthermore, the architecture incorporates "
        "an interactive multi-pane OpenCV dashboard, an Ollama LLM voice agent (Dr. Daniel Pauly), bilingual (English/German) "
        "live chat query handling, and automated multi-format report exports (CSV, PNG, Markdown, LaTeX/PDF, and DOCX)."
    )
    abs_p.paragraph_format.space_after = Pt(18)

    # Section 1: System Overview
    doc.add_heading("1. System Overview & Core Objectives", level=1)
    doc.add_paragraph(
        "The primary objective of the AquaPulse framework is to non-invasively monitor aquatic ecosystems in real time. "
        "Standard visual inspection of underwater video sources often suffers from lighting attenuation, turbidity, overlapping "
        "specimens, and rapid fish movement. The AquaPulse pipeline addresses these challenges through 6 core operational pillars:"
    )
    pillars = [
        "Detecting and classifying marine taxa frame-by-frame using CUDA-accelerated YOLO models.",
        "Assigning persistent track IDs to individual organisms across continuous frames using Kalman filtering and Re-ID features, effectively eliminating double-counting in biological censuses.",
        "Extracting high-resolution cropped specimen images for each unique track ID into a localized session repository.",
        "Forecasting unobserved apex predator populations (Y_n) from observed prey detections (X_n) using non-linear Lotka-Volterra dynamics and Euler-Maruyama stochastic differential equations.",
        "Quantifying instantaneous extinction risk probabilities P(Extinction) via Ensemble Kalman Filtering (EnKF).",
        "Providing an expert LLM dialogue interface (Dr. Daniel Pauly) with real-time bilingual query detection (English/German) and mouse-selected fish specimen telemetry context."
    ]
    for p_text in pillars:
        doc.add_paragraph(p_text, style='List Bullet')

    # Section 2: Modular Architecture
    doc.add_heading("2. Modular Codebase Architecture", level=1)
    doc.add_paragraph(
        "The codebase located at C:\\Users\\parsa\\Desktop\\Code\\3 - AI process is structured into high-cohesion, "
        "decoupled Python modules:"
    )

    modules_data = [
        ("mod_00_config_and_assets.py", "Hardware Acceleration & Asset Management",
         "Detects NVIDIA CUDA GPU hardware (e.g., NVIDIA GeForce RTX 4060 Laptop GPU) or defaults to CPU mode. Manages workspace asset discovery (johnny.gif) and creates unique, timestamped session directories (video_analysis_sessions/<VIDEO>_<TIMESTAMP>/) containing subfolders for csv/, plots/, analysis/, output/, and fish_images/."),
        
        ("mod_01_eco_census.py", "Biological Census Engine",
         "Tracks deduplicated unique individual specimen IDs per species using set registers. Calculates total specimen counts, biomass representation percentages, and outputs formatted fish_counts.csv reports with methodology statements."),
        
        ("mod_02_stats_and_plots.py", "Telemetry Visualization Suite",
         "Generates 20 high-resolution scientific plots summarizing session metrics. Includes population time series, extinction risk trajectories, phase space orbits, species abundance distributions, Shannon diversity (H'), Pielou evenness (J'), 2D spatial heatmaps, swimmer velocity histograms, Kalman Gain dynamics, and residual error diagnostics."),
        
        ("mod_03_vision_engine.py", "Computer Vision Reticle & FX Pipeline",
         "Renders customizable bounding box reticles, species color palettes, and motion vector trails. Implements Adaptive Underwater CLAHE (Contrast Limited Adaptive Histogram Equalization) image enhancement and synthetic water layer overlays."),
        
        ("mod_04_enkf_telemetry.py", "Stochastic Data Assimilation Engine",
         "Implements the 100-member Ensemble Kalman Filter. Integrates Lotka-Volterra predator-prey dynamics via Euler-Maruyama stochastic steps, computes empirical covariance matrices (C_n^(zy), C_n^(yy)), updates state vectors using empirical Kalman Gain (K_n), and calculates extinction probabilities."),
        
        ("mod_05_dialogue_and_ollama.py", "Ollama LLM & Audio Dialogue Interface",
         "Interfaces with the local Ollama LLM (llama3 / mistral). Powers the Dr. Daniel Pauly voice agent, GBIF species image fetcher, bilingual query engine (English/German auto-detection), and orchestrates executive Markdown, PDF, and DOCX exports."),
        
        ("mod_06_ui_dashboard.py", "4-Pane OpenCV Dashboard",
         "Renders an interactive UI divided into Pane 1 (Live Statistical Telemetry & EnKF Gauge), Pane 2 (Main Video Viewport), Pane 3 (Interactive Controls & Live Chat Portal), and Pane 4 (Comm Link, Selected Fish Specimen Reticle Zoom, and Relic Overlay)."),
        
        ("mod_07_pdf_exporter.py", "LaTeX PDF Generation Engine",
         "Parses report_template.tex, sanitizes input strings (escape_latex), populates species census data, constructs a 3-column LaTeX minipage specimen image gallery, embeds all 20 plots with analytical breakdowns, and compiles native PDFs using MiKTeX pdflatex."),
        
        ("manual_botsort.py", "Standalone BoT-SORT Module",
         "Provides standalone execution capabilities for BoT-SORT tracking with Camera Motion Compensation (GMC) and Re-ID embedding."),
        
        ("docx_report_generator.py", "Microsoft Word Report Exporter",
         "Generates structured ecological_analysis_report.docx documents containing complete hotkey registries, mathematical equations, and telemetry tables."),
        
        ("main.py", "System Orchestrator",
         "Initializes models, manages the video capture loop, processes mouse selection events, handles hotkeys (SPACE, S, F, V, R, X, L, A, W, M, T, C, J, TAB/CTRL+T/CTRL+C), and coordinates full session exports.")
    ]

    for fname, title, desc in modules_data:
        p = doc.add_paragraph()
        run_fn = p.add_run(f"• {fname} — {title}: ")
        run_fn.bold = True
        run_fn.font.color.rgb = RGBColor(0x00, 0x4B, 0x87)
        p.add_run(desc)

    # Section 3: Vision & Multi-Algorithm Tracking
    doc.add_heading("3. Computer Vision & Multi-Algorithm Tracking Framework", level=1)
    doc.add_paragraph(
        "The vision layer loads neural weights from the models/ directory in strict priority order:\n"
        "Priority Hierarchy: fish_model.pt -> best.pt -> medium.pt -> small.pt"
    )
    doc.add_paragraph(
        "Tracking is executed dynamically via Ultralytics YOLO with multi-algorithm support:"
    )
    doc.add_paragraph(
        "• BoT-SORT (botsort.yaml): Integrates Kalman filtering with Global Motion Compensation (GMC via sparse optical flow) and appearance Re-ID features to maintain track IDs during camera movement.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "• ByteTrack (bytetrack.yaml): High-speed association method that preserves low-confidence detection boxes to prevent track fragmentation.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "When a new track ID i is initialized, a high-resolution cropped image of the specimen is saved to:\n"
        "video_analysis_sessions/<SESSION_NAME>/fish_images/<SPECIES>.png (1 deduplicated image per species)."
    )

    # Section 4: Mathematical Framework
    doc.add_heading("4. Mathematical Framework & Data Assimilation", level=1)
    doc.add_paragraph(
        "The mathematical foundation models prey abundance X_n (observed via YOLO) and hidden apex predator density Y_n (unobserved latent state)."
    )

    doc.add_heading("4.1 Base Stochastic Differential Equation (SDE)", level=2)
    doc.add_paragraph(
        "Using the Euler-Maruyama numerical scheme with process noise:"
    )
    eq1 = doc.add_paragraph("Z_(n+1) = Z_n + Δt · f(Z_n) + √(Δt) · ζ_n,    ζ_n ~ N(0, Σ)")
    eq1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq1.runs[0].bold = True

    doc.add_heading("4.2 Ecosystem Lotka-Volterra Formulation", level=2)
    doc.add_paragraph("Applying stochastic forcing to predator-prey dynamics:")
    doc.add_paragraph("1. Prey Evolution (Observed Vision State X):")
    eq_prey = doc.add_paragraph("X_(n+1) = X_n + Δt · (α·X_n - β·X_n·Y_n) + √(Δt) · σ_X · X_n · ζ_n^X")
    eq_prey.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq_prey.runs[0].bold = True

    doc.add_paragraph("2. Predator Evolution (Hidden Latent Variable Y):")
    eq_pred = doc.add_paragraph("Y_(n+1) = Y_n + Δt · (δ·X_n·Y_n - γ·Y_n) + √(Δt) · σ_Y · Y_n · ζ_n^Y")
    eq_pred.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq_pred.runs[0].bold = True

    doc.add_paragraph(
        "where α is prey birth rate, β is predation rate, δ is predator growth efficiency, γ is predator mortality rate, "
        "and σ_X, σ_Y represent environmental volatility."
    )

    doc.add_heading("4.3 Ensemble Kalman Filter (EnKF) Algorithm", level=2)
    doc.add_paragraph("Following Sprungk (2023), the EnKF operates with M = 100 state ensemble members:")

    enkf_steps = [
        ("1. Initialization Phase",
         "Draw M initial samples z_0^(m) = [X_0^(m), Y_0^(m)]^T, m = 1,...,M i.i.d. from prior distribution π_0, and set the initial ensemble state set:\n"
         "𝔨_0^a = { z_0^(1), ..., z_0^(M) }"),
        
        ("2. Sequential Forecasting Phase (For n = 1, 2, 3, ...)",
         "Sample independently for each ensemble member m = 1,...,M:\n"
         "z_n^(m) ~ P(z_(n-1)^(m)),    y_n^(m) ~ N(H(z_(n-1)^(m)), Γ)\n"
         "where H = [1, 0] is the observation operator extracting prey counts, Γ is measurement error variance, and set the forecast ensemble state set:\n"
         "𝔨_n^f = { z_n^(1), ..., z_n^(M) }"),
        
        ("3. Empirical Covariance Computation",
         "Compute empirical ensemble state mean z̄_n and empirical measurement mean ȳ_n:\n"
         "z̄_n = (1/M) Σ z_n^(m),    ȳ_n = (1/M) Σ y_n^(m)\n\n"
         "Calculate empirical cross-covariance C_n^(zy) and measurement error covariance C_n^(yy):\n"
         "C_n^(zy) = (1 / (M-1)) Σ [ z_n^(m) - z̄_n ] [ y_n^(m) - ȳ_n ]^T\n"
         "C_n^(yy) = (1 / (M-1)) Σ [ y_n^(m) - ȳ_n ] [ y_n^(m) - ȳ_n ]^T"),
        
        ("4. Empirical Kalman Gain & Analysis Update",
         "Compute the empirical Kalman Gain matrix K_n:\n"
         "K_n := C_n^(zy) · (C_n^(yy))^(-1)\n\n"
         "Update each ensemble state vector using live YOLO measurement observation y_n:\n"
         "z_n^(m) = z_n^(m) + K_n · [ y_n - y_n^(m) ]\n"
         "and set the updated analysis ensemble state set:\n"
         "𝔨_n^a = { z_n^(1), ..., z_n^(M) }"),
        
        ("5. Extinction Risk Probability Evaluation",
         "Evaluate the proportion of ensemble members crossing critical threshold X_crit = 2.0:\n"
         "P(Extinction) = (1/M) Σ 𝕀( X_n^(m) ≤ X_crit )")
    ]

    for step_title, step_body in enkf_steps:
        p_st = doc.add_paragraph()
        run_st = p_st.add_run(step_title)
        run_st.bold = True
        run_st.font.color.rgb = RGBColor(0x00, 0x4B, 0x87)
        
        p_sb = doc.add_paragraph(step_body)
        p_sb.paragraph_format.left_indent = Inches(0.25)
        p_sb.paragraph_format.space_after = Pt(12)

    # Section 5: Interactive UI
    doc.add_heading("5. Interactive UI & Bilingual Language Engine", level=1)
    ui_features = [
        ("Target Selection via Mouse", "Clicking a bounding box in the main video pane locks the target specimen (locked_target). The reticle highlights the target and displays a high-resolution zoomed preview in Pane 4."),
        ("Targeted Dr. Pauly Calls", "When Dr. Pauly is called (via key C, button [C] DR. PAULY, or chat), the system passes the mouse-selected specimen's ID, species name, and confidence score into the LLM context."),
        ("Bilingual Live Chat Portal", "Opened via button [CHAT] [TAB / CTRL+T / CTRL+C] or hotkeys (TAB, CTRL+T, CTRL+C). The prompt engine automatically detects whether the user's question is typed in English or German and responds in the matching language."),
        ("Pause Inspection Mode", "Redundant model tracking is bypassed during pause, allowing full mouse target selection, UI control toggling, and Dr. Pauly calls on static frames without CPU/GPU load.")
    ]
    for title, desc in ui_features:
        p = doc.add_paragraph()
        r = p.add_run(f"• {title}: ")
        r.bold = True
        p.add_run(desc)

    # Section 6: Export Artifacts
    doc.add_heading("6. Session Output Artifacts", level=1)
    doc.add_paragraph(
        "At the end of each session (or upon closing the application), the pipeline automatically compiles 7 comprehensive export artifacts:"
    )
    artifacts = [
        "1. Tracked MP4 Video: Saved to output/tracked_<VIDEO>.mp4",
        "2. Deduplicated CSV Report: Saved to csv/fish_counts.csv",
        "3. 20 Scientific Telemetry Plots: Saved to plots/01_...png through plots/20_...png",
        "4. Cropped Fish Specimen Gallery: Saved to fish_images/<SPECIES>.png (1 image per species)",
        "5. Executive Markdown Report: Saved to analysis/ollama_marine_report.md",
        "6. LaTeX Source & Compiled Native PDF: Saved to analysis/ollama_marine_report.tex & .pdf",
        "7. Microsoft Word DOCX Document: Saved to analysis/ecological_analysis_report.docx"
    ]
    for art in artifacts:
        doc.add_paragraph(art, style='List Bullet')

    # Section 7: References
    doc.add_heading("7. References", level=1)
    ref_p = doc.add_paragraph(
        "[1] Sprungk, B. (2023). Probabilistic Forecasting and Data Assimilation. Lecture Course Notes, TU Bergakademie Freiberg (TUBAF)."
    )
    ref_p.paragraph_format.left_indent = Inches(0.25)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Successfully generated Word DOCX documentation at: {output_path}")

if __name__ == "__main__":
    doc_dir = r"C:\Users\parsa\Desktop\Code\4 - Documentation"
    out_docx = os.path.join(doc_dir, "aquapulse_system_architecture_documentation.docx")
    create_system_architecture_docx(out_docx)
